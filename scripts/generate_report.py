from pathlib import Path
import subprocess
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "entrega-fullstack-riquelme.docx"
PDF_OUTPUT = ROOT / "entrega-fullstack-riquelme.pdf"
CODE_FILES = [
    "package.json",
    "server.js",
    "public/index.html",
    "public/style.css",
    "public/script.js",
    "README.md",
]


def set_margins(section):
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(26)
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(3)

    for name, size, color in [
        ("Heading 1", 20, RGBColor(0, 0, 0)),
        ("Heading 2", 16, RGBColor(0, 0, 0)),
        ("Heading 3", 14, RGBColor(67, 67, 67)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    code = styles.add_style("Code Block", 1)
    code.font.name = "Courier New"
    code.font.size = Pt(8)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.15)


def add_kv_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    return table


def add_cover(document):
    document.add_paragraph("Trabalho Academico", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Praticas 1 e 2 - Desenvolvimento Full Stack e Versionamento Git")
    run.font.size = Pt(14)

    document.add_paragraph()
    add_kv_table(
        document,
        [
            ("Aluno", "Riquelme"),
            ("Tema", "Aplicacao web com HTML, JavaScript e Node.js"),
            ("Entrega", "Codigo no GitHub e relatorio em PDF"),
            ("Repositorio local", str(ROOT)),
        ],
    )

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("2026").bold = True
    document.add_page_break()


def add_intro(document):
    document.add_heading("Resumo da atividade", level=1)
    document.add_paragraph(
        "Este trabalho apresenta uma aplicacao Full Stack simples que demonstra a comunicacao "
        "entre cliente e servidor. O usuario digita um nome no front-end, o navegador envia uma "
        "requisicao HTTP para o back-end em Node.js e o servidor retorna uma saudacao personalizada "
        "em formato JSON."
    )
    document.add_paragraph(
        "A segunda parte da pratica foi aplicada com Git, criando commits, uma branch de feature, "
        "uma branch de hotfix e merges para integrar as alteracoes na branch principal."
    )

    document.add_heading("Tecnologias utilizadas", level=2)
    for item in ["HTML", "CSS", "JavaScript", "Node.js", "Git", "GitHub"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Como executar o projeto", level=2)
    for step in [
        "Acessar a pasta do projeto.",
        "Executar o comando npm start.",
        "Abrir http://localhost:3000 no navegador.",
        "Digitar um nome e enviar o formulario.",
    ]:
        document.add_paragraph(step, style="List Number")


def add_git_history(document):
    document.add_heading("Fluxo de versionamento", level=1)
    document.add_paragraph(
        "O historico Git foi organizado para representar um fluxo comum de desenvolvimento: "
        "a branch main recebeu a estrutura inicial, a branch de feature adicionou a interface "
        "e a branch de hotfix corrigiu uma validacao do servidor."
    )

    log = subprocess.check_output(
        ["git", "log", "--oneline", "--graph", "--all", "--decorate", "-n", "8"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
    )
    document.add_heading("Historico Git", level=2)
    for line in log.splitlines():
        paragraph = document.add_paragraph(style="Code Block")
        paragraph.add_run(line)


def add_code_appendix(document):
    document.add_heading("Anexos - Codigo-fonte", level=1)
    document.add_paragraph(
        "Os arquivos abaixo compoem a aplicacao entregue no repositorio."
    )

    for relative_path in CODE_FILES:
        file_path = ROOT / relative_path
        document.add_heading(relative_path, level=2)
        content = file_path.read_text(encoding="utf-8")
        for line in content.splitlines():
            paragraph = document.add_paragraph(style="Code Block")
            paragraph.add_run(line if line else " ")


def main():
    document = Document()
    for section in document.sections:
        set_margins(section)
    configure_styles(document)

    add_cover(document)
    add_intro(document)
    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_git_history(document)
    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_code_appendix(document)

    document.save(OUTPUT)
    build_pdf()
    print(OUTPUT)
    print(PDF_OUTPUT)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName="Helvetica",
            fontSize=26,
            leading=32,
            alignment=1,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1",
            parent=styles["Heading1"],
            fontName="Helvetica",
            fontSize=18,
            leading=23,
            textColor=colors.black,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2",
            parent=styles["Heading2"],
            fontName="Helvetica",
            fontSize=14,
            leading=18,
            textColor=colors.black,
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PDFCode",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7.2,
            leading=8.6,
            leftIndent=10,
            firstLineIndent=0,
            spaceAfter=0,
            splitLongWords=True,
            wordWrap="CJK",
        )
    )

    story = []
    story.append(Paragraph("Trabalho Academico", styles["CoverTitle"]))
    story.append(
        Paragraph(
            "Praticas 1 e 2 - Desenvolvimento Full Stack e Versionamento Git",
            styles["Body"],
        )
    )
    story.append(Spacer(1, 0.35 * inch))
    table = Table(
        [
            ["Aluno", "Riquelme"],
            ["Tema", "Aplicacao web com HTML, JavaScript e Node.js"],
            ["Entrega", "Codigo no GitHub e relatorio em PDF"],
            ["Repositorio local", str(ROOT)],
        ],
        colWidths=[1.45 * inch, 4.85 * inch],
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.35 * inch))
    story.append(Paragraph("2026", styles["Body"]))
    story.append(PageBreak())

    story.append(Paragraph("Resumo da atividade", styles["H1"]))
    story.append(
        Paragraph(
            "Este trabalho apresenta uma aplicacao Full Stack simples que demonstra a comunicacao "
            "entre cliente e servidor. O usuario digita um nome no front-end, o navegador envia uma "
            "requisicao HTTP para o back-end em Node.js e o servidor retorna uma saudacao personalizada "
            "em formato JSON.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "A segunda parte da pratica foi aplicada com Git, criando commits, uma branch de feature, "
            "uma branch de hotfix e merges para integrar as alteracoes na branch principal.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Tecnologias utilizadas", styles["H2"]))
    for item in ["HTML", "CSS", "JavaScript", "Node.js", "Git", "GitHub"]:
        story.append(Paragraph(f"- {item}", styles["Body"]))

    story.append(Paragraph("Como executar o projeto", styles["H2"]))
    for index, step in enumerate(
        [
            "Acessar a pasta do projeto.",
            "Executar o comando npm start.",
            "Abrir http://localhost:3000 no navegador.",
            "Digitar um nome e enviar o formulario.",
        ],
        start=1,
    ):
        story.append(Paragraph(f"{index}. {step}", styles["Body"]))

    story.append(PageBreak())
    story.append(Paragraph("Fluxo de versionamento", styles["H1"]))
    story.append(
        Paragraph(
            "O historico Git foi organizado para representar um fluxo comum de desenvolvimento: "
            "a branch main recebeu a estrutura inicial, a branch de feature adicionou a interface "
            "e a branch de hotfix corrigiu uma validacao do servidor.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Historico Git", styles["H2"]))
    log = subprocess.check_output(
        ["git", "log", "--oneline", "--graph", "--all", "--decorate", "-n", "8"],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
    )
    for line in log.splitlines():
        story.append(Paragraph(escape(line), styles["PDFCode"]))

    story.append(PageBreak())
    story.append(Paragraph("Anexos - Codigo-fonte", styles["H1"]))
    story.append(Paragraph("Os arquivos abaixo compoem a aplicacao entregue no repositorio.", styles["Body"]))

    for relative_path in CODE_FILES:
        story.append(Paragraph(relative_path, styles["H2"]))
        for line in (ROOT / relative_path).read_text(encoding="utf-8").splitlines():
            safe_line = escape(line).replace(" ", "&nbsp;")
            story.append(Paragraph(safe_line or "&nbsp;", styles["PDFCode"]))

    pdf = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="Trabalho Full Stack - Riquelme",
        author="Riquelme",
    )
    pdf.build(story)


if __name__ == "__main__":
    main()
